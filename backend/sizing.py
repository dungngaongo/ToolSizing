import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
from datetime import datetime
import math
import requests
import base64
import json

# --- CẤU HÌNH HỆ THỐNG ---
BACKEND_API_URL = "http://localhost:8081/api"

# --- HÀM HỖ TRỢ LƯU DỮ LIỆU ---
def get_image_base64(uploaded_file):
    """Chuyển đổi ảnh upload thành base64 string"""
    if uploaded_file is not None:
        bytes_data = uploaded_file.getvalue()
        return base64.b64encode(bytes_data).decode('utf-8')
    return None

def get_existing_dinh_co_content(project_id):
    """Lấy dữ liệu dinhCoHeThongContent hiện có từ database"""
    try:
        response = requests.get(f"{BACKEND_API_URL}/project-data/project/{project_id}")
        if response.status_code == 200:
            data = response.json()
            content = data.get("dinhCoHeThongContent")
            if content:
                return json.loads(content)
        return {}
    except Exception as e:
        st.error(f"Lỗi khi lấy dữ liệu: {e}")
        return {}

def save_dinh_co_module(project_id, module_name, module_data):
    """Lưu dữ liệu một module vào dinhCoHeThongContent"""
    try:
        # Lấy dữ liệu hiện có
        existing_content = get_existing_dinh_co_content(project_id)
        
        # Cập nhật module mới
        existing_content[module_name] = module_data
        
        # Gửi lên API
        payload = {
            "dinhCoHeThongContent": json.dumps(existing_content, ensure_ascii=False)
        }
        
        response = requests.put(
            f"{BACKEND_API_URL}/project-data/project/{project_id}",
            json=payload,
            headers={"Content-Type": "application/json"}
        )
        
        if response.status_code == 200:
            return True, "Lưu thành công!"
        else:
            return False, f"Lỗi API: {response.text}"
    except Exception as e:
        return False, str(e)

def save_redis_sizing(system_info_id, redis_data):
    """Gửi dữ liệu sizing Redis xuống Backend API"""
    try:
        url = f"{BACKEND_API_URL}/redis/system-info/{system_info_id}"
        headers = {"Content-Type": "application/json"}
        response = requests.post(url, json=redis_data, headers=headers)
        if response.status_code == 200:
            return True, response.json()
        else:
            return False, response.text
    except Exception as e:
        return False, str(e)

def load_saved_data_from_db(project_id):
    """Load dữ liệu đã lưu từ database vào session_state"""
    if not project_id:
        return
    
    # Chỉ load 1 lần khi khởi động
    if "data_loaded" in st.session_state and st.session_state.data_loaded == project_id:
        return
    
    try:
        content = get_existing_dinh_co_content(project_id)
        if content:
            # Load Oracle RAC data
            if "oracleRac_module" in content:
                oracle_data = content["oracleRac_module"]
                st.session_state.oracle_saved_data = oracle_data
                # Load servers từ database
                if "server" in oracle_data and oracle_data["server"]:
                    st.session_state.servers = []
                    for s in oracle_data["server"]:
                        st.session_state.servers.append({
                            "ip": s.get("ip", ""),
                            "vcpu": s.get("vCpu", 32),
                            "ram": s.get("ram", 64),
                            "cpu_load": s.get("vCpu_percent", 30),
                            "ram_load": s.get("ram_percent", 80)
                        })
            
            # Load Redis data
            if "redis_module" in content:
                st.session_state.redis_saved_data = content["redis_module"]
            
            # Load MongoDB data
            if "mongo_module" in content:
                st.session_state.mongo_saved_data = content["mongo_module"]
            
            # Load PostgreSQL data
            if "postgresql_module" in content:
                st.session_state.pg_saved_data = content["postgresql_module"]
            
            st.session_state.data_loaded = project_id
    except Exception as e:
        st.warning(f"Không thể load dữ liệu: {e}")

# --- THIẾT LẬP TRANG ---
st.set_page_config(page_title="Cong cu Dinh co Ha tang", layout="wide")

# --- CSS TÙY CHỈNH ---
st.markdown("""
<style>
    /* Font chữ toàn hệ thống */
    html, body, [class*="css"] {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    
    /* Màu nền chính */
    .main { background-color: #f4f6f9; }
    
    /* Header Bar */
    .header-style {
        background-color: #2c3e50;
        padding: 20px;
        border-radius: 8px;
        color: white;
        margin-bottom: 25px;
        border-left: 5px solid #e74c3c;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .header-title {
        font-size: 24px;
        font-weight: 600;
        margin: 0;
        text-transform: uppercase;
    }
    .header-subtitle {
        font-size: 14px;
        opacity: 0.8;
        margin-top: 5px;
    }

    /* Card/Container Styling */
    div[data-testid="stVerticalBlockBorderWrapper"] {
        background-color: white;
        border-radius: 8px;
        padding: 15px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        border: 1px solid #e0e0e0;
        margin-bottom: 15px;
    }

    /* Metric Box Styling */
    div[data-testid="stMetric"] {
        background-color: #ffffff;
        border: 1px solid #dcdcdc;
        padding: 15px;
        border-radius: 6px;
        border-left: 4px solid #3498db;
    }
    div[data-testid="stMetric"] label {
        font-weight: bold;
        color: #555;
    }

    /* Button Styling */
    .stButton>button {
        background-color: #2980b9;
        color: white;
        font-weight: 600;
        border-radius: 4px;
        border: none;
        height: 45px;
        width: 100%;
        transition: background 0.3s;
    }
    .stButton>button:hover {
        background-color: #1c5980;
    }

    /* Slider Styling (Màu đỏ theo yêu cầu) */
    div.stSlider > div[data-baseweb = "slider"] > div > div > div[role="slider"]{
        background-color: #ff4b4b !important;
        box-shadow: 0 0 5px rgba(255, 75, 75, 0.5);
    }
    div.stSlider > div[data-baseweb = "slider"] > div > div > div > div {
        background-color: #ff4b4b !important;
    }
    
    /* Input field adjustment */
    .stTextInput input, .stNumberInput input {
        background-color: #f8f9fa;
    }
</style>
""", unsafe_allow_html=True)

# --- KHỞI TẠO SESSION STATE (CHO ORACLE/BASELINE) ---
if "servers" not in st.session_state:
    st.session_state.servers = [
        {"ip": "10.207.252.1", "vcpu": 32, "ram": 64, "cpu_load": 30, "ram_load": 80}
    ]

def add_server():
    idx = len(st.session_state.servers) + 1
    st.session_state.servers.append({
        "ip": f"10.207.252.{idx}", 
        "vcpu": 32, 
        "ram": 64, 
        "cpu_load": 30, 
        "ram_load": 80
    })

def remove_server(index):
    if len(st.session_state.servers) > 0:
        st.session_state.servers.pop(index)

# --- TIÊU ĐỀ TRANG ---
st.markdown("""
<div class="header-style">
    <div class="header-title">Công cụ Định cỡ Hạ tầng Cơ sở dữ liệu</div>
    <div class="header-subtitle">Hỗ trợ lập kế hoạch tài nguyên hệ thống (Sizing Tool)</div>
</div>
""", unsafe_allow_html=True)

# --- THANH BÊN (SIDEBAR) ---
with st.sidebar:
    st.header("Cấu hình chung")
    
    # Nhập Project ID
    query_params = st.query_params
    url_project_id = query_params.get("projectId", "")
    
    if url_project_id:
        project_id = url_project_id
        st.info(f"Mã dự án: {project_id}")
    else:
        project_id = st.text_input("Mã dự án (Project ID)", placeholder="Nhập ID...")

    # Load dữ liệu đã lưu từ database
    if project_id:
        load_saved_data_from_db(project_id)

    st.markdown("---")
    
    # Chọn loại DB
    st.subheader("Hệ thống Đích")
    db_type = st.selectbox(
        "Loại Cơ sở dữ liệu",
        ["Oracle RAC", "MariaDB MaxScale", "PostgreSQL", "Redis", "MongoDB"]
    )

    st.markdown("---")
    
    # Tham số an toàn
    with st.expander("Tham số An toàn & Tải", expanded=True):
        growth_factor = st.slider("Hệ số dự phòng", 1.0, 2.0, 1.1, 0.1, help="Dự phòng cho tăng trưởng dữ liệu")
        cpu_threshold = st.slider("Ngưỡng tải CPU tối đa", 0.5, 1.0, 0.75, 0.05)
        ram_threshold = st.slider("Ngưỡng tải RAM tối đa", 0.5, 1.0, 0.90, 0.05)
        disk_threshold = st.slider("Ngưỡng đầy Disk tối đa", 0.5, 1.0, 0.80, 0.05)

# ==========================================
# LOGIC XỬ LÝ CHÍNH
# ==========================================

# --- TRƯỜNG HỢP 1: REDIS SIZING ---
if db_type == "Redis":
    
    # Load dữ liệu đã lưu
    redis_saved = st.session_state.get("redis_saved_data", {})
    redis_dauVao = redis_saved.get("dauVao", {})
    redis_result = redis_saved.get("result", {})
    redis_so_lieu = redis_saved.get("so_lieu", [{}])[0] if redis_saved.get("so_lieu") else {}
    
    col_input, col_result = st.columns([1.2, 0.8], gap="medium")
    
    with col_input:
        st.subheader("Tham số Đầu vào")
        
        tab_info, tab_sizing = st.tabs(["Thông tin Mô hình", "Số liệu Tính toán"])
        
        with tab_info:
            with st.container(border=True):
                st.markdown("**1. Sơ đồ Logic**")
                
                col_upload, col_icon = st.columns([5, 1])
                
                with col_upload:
                    uploaded_logic_img = st.file_uploader("Tải lên sơ đồ", type=['png', 'jpg'], key="redis_logic", label_visibility="collapsed")
                
                with col_icon:
                    with st.popover("❓", help="Xem ảnh minh họa"):
                        st.image("mo-hinh-he-thong/redis_example.png", use_container_width=True, caption="Ví dụ sơ đồ Logic Redis")
                
                if uploaded_logic_img:
                    st.image(uploaded_logic_img, use_container_width=True, caption="Sơ đồ Logic")

            with st.container(border=True):
                st.markdown("**3. Thông tin Nghiệp vụ**")
                redis_desc = st.text_area("Mô tả module", height=100, value=redis_dauVao.get("mo_ta", ""), placeholder="Mô tả chức năng chính của module...")
                redis_purpose = st.text_area("Mục đích sử dụng", height=100, value=redis_dauVao.get("muc_dich", ""), placeholder="Ví dụ: Cache, Session, Pub/Sub...")

        with tab_sizing:
            method = st.radio("Phương pháp tính toán", ["Theo số lượng Key (Khuyên dùng)", "Tuyến tính theo cấu hình cũ"], horizontal=True)
            
            redis_params = {}

            if "Key" in method:
                with st.container(border=True):
                    st.markdown("#### Chỉ số Key")
                    c_key1, c_key2 = st.columns(2)
                    
                    with c_key1:
                        key_count = st.number_input("Tổng số lượng Key (A)", min_value=1, value=redis_so_lieu.get("key", 1000000), step=100000, format="%d")
                        st.caption("Gợi ý: Dùng lệnh dbsize hoặc info keyspace")
                        uploaded_key_proof = st.file_uploader("Ảnh sở cứ (Số lượng Key)", type=['png', 'jpg'], key="kp")
                    
                    with c_key2:
                        avg_size_kb = st.number_input("Kích thước TB/Key (KB) (B)", min_value=0.0, value=float(redis_so_lieu.get("kich_thuoc", 2.0)), step=0.1)
                        st.caption("Gợi ý: Dùng lệnh memory usage")
                        uploaded_size_proof = st.file_uploader("Ảnh sở cứ (Kích thước)", type=['png', 'jpg'], key="sp")

                with st.container(border=True):
                    st.markdown("#### Chiến lược Cluster")
                    c_clus1, c_clus2 = st.columns(2)
                    with c_clus1:
                        system_criticality = st.selectbox("Mức độ quan trọng", ["Thường (1 Master - 1 Slave)", "Đặc biệt quan trọng (1 Master - 2 Slave)"])
                    with c_clus2:
                        num_shards = st.number_input("Số Shard (Master) dự kiến (N)", min_value=1, value=3, step=2, help="Nên là số lẻ")

                redis_params = {
                    "method": "keys", "A": key_count, "B_KB": avg_size_kb,
                    "criticality": system_criticality, "N": num_shards,
                    "proof_key": uploaded_key_proof, "proof_size": uploaded_size_proof
                }
            
            else: # Phương pháp Tuyến tính
                with st.container(border=True):
                    st.warning("Lưu ý: Phương pháp này chỉ nhân bản cấu hình cũ, không tối ưu cho Redis.")
                    c1, c2, c3 = st.columns(3)
                    curr_ram = c1.number_input("RAM Hiện tại (GB)", value=16)
                    curr_cpu = c2.number_input("vCPU Hiện tại", value=4)
                    curr_disk = c3.number_input("Disk Hiện tại (GB)", value=50)
                    
                    l1, l2 = st.columns(2)
                    load_ram = l1.slider("% RAM đang dùng", 0, 100, 70)
                    load_cpu = l2.slider("% CPU đang dùng", 0, 100, 30)
                    target_scale = st.number_input("Tỉ lệ mở rộng (Scale Ratio)", value=1.5)

                    redis_params = {
                        "method": "linear", "curr_ram": curr_ram, "curr_cpu": curr_cpu, 
                        "curr_disk": curr_disk, "load_ram": load_ram, "load_cpu": load_cpu,
                        "scale": target_scale
                    }

    with col_result:
        st.subheader("Kết quả Định cỡ")
        
        # Hiển thị kết quả đã lưu nếu có
        if redis_result:
            st.info("📋 Đã tìm thấy kết quả đã lưu trước đó")
            with st.container(border=True):
                st.markdown(f"**Kiến trúc:** `{redis_result.get('kien_truc', 'N/A')}`")
                st.markdown(f"**Số Node:** {redis_result.get('node', 'N/A')}")
                m1, m2 = st.columns(2)
                m1.metric("vCPU", f"{redis_result.get('vcpu', 'N/A')} Core")
                m2.metric("RAM", f"{redis_result.get('ram', 'N/A')} GB")
            st.divider()
        
        if st.button("TÍNH TOÁN TÀI NGUYÊN REDIS", type="primary", use_container_width=True):
            
            final_model = ""
            node_config = {}
            total_size_gb = 0

            if redis_params["method"] == "keys":
                total_size_gb = (redis_params["A"] * redis_params["B_KB"]) / (1024 * 1024)
                
                if total_size_gb < 32:
                    final_model = "Redis Sentinel"
                    req_ram = (total_size_gb * 1.1) / 0.8
                    node_config = {"qty": 3, "vCPU": 16, "RAM": math.ceil(req_ram), "Disk": math.ceil(req_ram * 4)}
                    st.info(f"Tổng dung lượng Data: {total_size_gb:.2f} GB (<32GB) -> Đề xuất mô hình Sentinel")
                else:
                    final_model = "Redis Cluster"
                    N = redis_params["N"]
                    slaves = 2 if "Đặc biệt" in redis_params["criticality"] else 1
                    total_nodes = N * (1 + slaves)
                    req_ram = (total_size_gb * 1.1) / 0.8 / N
                    node_config = {"qty": total_nodes, "vCPU": 8, "RAM": math.ceil(req_ram), "Disk": math.ceil(req_ram * 4)}
                    st.info(f"Tổng dung lượng Data: {total_size_gb:.2f} GB (>32GB) -> Đề xuất mô hình Cluster ({N} Shards)")
            
            else: 
                p = redis_params
                req_ram_total = (p["curr_ram"] * (p["load_ram"]/100) * p["scale"] * growth_factor) / 0.9
                req_cpu_total = (p["curr_cpu"] * (p["load_cpu"]/100) * p["scale"] * growth_factor) / 0.75
                final_model = "Redis Sentinel (Tuyến tính)"
                node_config = {"qty": 3, "vCPU": max(4, math.ceil(req_cpu_total/3)), "RAM": math.ceil(req_ram_total/3), "Disk": math.ceil(req_ram_total/3)*4}

            # Lưu kết quả vào session_state
            st.session_state.redis_result = {
                "final_model": final_model,
                "node_config": node_config,
                "redis_params": redis_params,
                "redis_desc": redis_desc,
                "redis_purpose": redis_purpose,
                "uploaded_logic_img": uploaded_logic_img
            }

            with st.container(border=True):
                st.markdown(f"**Kiến trúc đề xuất:** `{final_model}`")
                st.markdown(f"**Tổng số Node:** {node_config['qty']}")
                
                st.divider()
                st.markdown("#### Cấu hình mỗi Node")
                
                m1, m2 = st.columns(2)
                m1.metric("vCPU", f"{node_config['vCPU']} Core")
                m2.metric("RAM", f"{node_config['RAM']} GB", delta="Đã bao gồm Buffer")
                
                m3, m4 = st.columns(2)
                m3.metric("Lưu trữ (Disk)", f"{node_config['Disk']} GB", "Tỉ lệ 4x RAM")
                m4.metric("Mạng (Network)", "10 Gbps")

            st.divider()
            doc = Document()
            doc.add_heading('BAO CAO DINH CO REDIS', 0)
            doc.add_paragraph(f"Mo hinh: {final_model}")
            doc.add_paragraph(f"Cau hinh Node: {node_config['vCPU']} vCPU, {node_config['RAM']} GB RAM")
            b = BytesIO()
            doc.save(b)
            b.seek(0)
            st.download_button("Tải báo cáo (.docx)", data=b, file_name="redis_sizing.docx", use_container_width=True)

        # Nút lưu đặt bên ngoài if button để luôn hiển thị khi có kết quả
        if "redis_result" in st.session_state:
            result = st.session_state.redis_result
            
            if st.button("💾 Lưu kết quả Redis vào Database", use_container_width=True, key="save_redis", type="primary"):
                if project_id:
                    redis_params = result["redis_params"]
                    node_config = result["node_config"]
                    final_model = result["final_model"]
                    
                    # Chuẩn bị dữ liệu theo cấu trúc mới
                    redis_module_data = {
                        "dauVao": {
                            "logic_image": get_image_base64(result.get("uploaded_logic_img")),
                            "mo_ta": result.get("redis_desc", ""),
                            "muc_dich": result.get("redis_purpose", "")
                        },
                        "phuong_phap": redis_params.get("method", ""),
                        "so_lieu": []
                    }
                    
                    # Thêm số liệu theo phương pháp
                    if redis_params.get("method") == "keys":
                        redis_module_data["so_lieu"].append({
                            "key": redis_params.get("A", 0),
                            "kich_thuoc": redis_params.get("B_KB", 0),
                            "key_image": get_image_base64(redis_params.get("proof_key")),
                            "kich_thuoc_image": get_image_base64(redis_params.get("proof_size")),
                            "importance": redis_params.get("criticality", ""),
                            "shard": redis_params.get("N", 0)
                        })
                    else:
                        redis_module_data["so_lieu"].append({
                            "ram": redis_params.get("curr_ram", 0),
                            "ram_percent": redis_params.get("load_ram", 0),
                            "vcpu": redis_params.get("curr_cpu", 0),
                            "vcpu_percent": redis_params.get("load_cpu", 0),
                            "disk": redis_params.get("curr_disk", 0),
                            "scale_ratio": redis_params.get("scale", 1.0)
                        })
                    
                    # Thêm kết quả
                    redis_module_data["result"] = {
                        "kien_truc": final_model,
                        "node": node_config.get("qty", 0),
                        "vcpu": node_config.get("vCPU", 0),
                        "ram": node_config.get("RAM", 0),
                        "disk": node_config.get("Disk", 0),
                        "network": "10 Gbps"
                    }
                    
                    success, msg = save_dinh_co_module(project_id, "redis_module", redis_module_data)
                    if success:
                        st.success("✅ Lưu dữ liệu Redis thành công!")
                    else:
                        st.error(f"❌ Lỗi: {msg}")
                else:
                    st.error("Chưa nhập Mã dự án (Project ID)")

# --- TRƯỜNG HỢP 2: MONGODB SIZING ---
elif db_type == "MongoDB":
    
    # Load dữ liệu đã lưu
    mongo_saved = st.session_state.get("mongo_saved_data", {})
    mongo_dauVao = mongo_saved.get("dauVao", {})
    mongo_result = mongo_saved.get("result", {})
    mongo_so_lieu = mongo_saved.get("so_lieu", [{}])[0] if mongo_saved.get("so_lieu") else {}
    
    col_input, col_result = st.columns([1.2, 0.8], gap="medium")
    
    with col_input:
        st.subheader("Tham số Đầu vào (MongoDB)")
        
        tab_info, tab_sizing = st.tabs(["Thông tin Mô hình", "Số liệu Tính toán"])
        
        with tab_info:
            with st.container(border=True):
                st.markdown("**1. Sơ đồ Logic**")
                
                col_upload, col_icon = st.columns([5, 1])
                
                with col_upload:
                    uploaded_logic_img = st.file_uploader("Tải lên sơ đồ", type=['png', 'jpg'], key="mongo_logic", label_visibility="collapsed")
                
                with col_icon:
                    with st.popover("❓", help="Xem ảnh minh họa"):
                        st.image("mo-hinh-he-thong/mongodb_example.png", use_container_width=True, caption="Ví dụ sơ đồ Logic MongoDB")
                
                if uploaded_logic_img:
                    st.image(uploaded_logic_img, use_container_width=True, caption="Sơ đồ Logic")

            with st.container(border=True):
                st.markdown("**3. Thông tin Nghiệp vụ**")
                mongo_desc = st.text_area("Mô tả module", height=100, value=mongo_dauVao.get("mo_ta", ""), placeholder="Chức năng chính, loại dữ liệu lưu trữ...", key="mongo_desc")
                mongo_purpose = st.text_area("Mục đích sử dụng", height=100, value=mongo_dauVao.get("muc_dich", ""), placeholder="Ví dụ: Lưu trữ Log, Catalog sản phẩm...", key="mongo_purp")

        with tab_sizing:
            method = st.radio("Phương pháp tính toán", ["Theo số lượng Document (Khuyên dùng)", "Tuyến tính theo cấu hình cũ"], horizontal=True, key="mongo_method")
            
            mongo_params = {}

            if "Document" in method:
                with st.container(border=True):
                    st.markdown("#### Chỉ số Document (Bản ghi)")
                    c_doc1, c_doc2 = st.columns(2)
                    
                    with c_doc1:
                        doc_count = st.number_input("Tổng số lượng Document", min_value=1, value=mongo_so_lieu.get("document", 5000000), step=100000, format="%d", help="Tổng số bản ghi dự kiến lưu trữ")
                        st.caption("Nguồn: Lệnh `db.collection.count()`")
                        uploaded_doc_proof = st.file_uploader("Ảnh sở cứ (Số lượng)", type=['png', 'jpg'], key="kp_mongo")
                    
                    with c_doc2:
                        avg_doc_size_kb = st.number_input("Kích thước TB/Doc (KB)", min_value=0.0, value=float(mongo_so_lieu.get("kich_thuoc", 4.0)), step=0.1)
                        st.caption("Nguồn: Lệnh `db.collection.stats()`")
                        uploaded_size_proof = st.file_uploader("Ảnh sở cứ (Kích thước)", type=['png', 'jpg'], key="sp_mongo")

                with st.container(border=True):
                    st.markdown("#### Tỷ lệ Dữ liệu & Index")
                    c_ws1, c_ws2 = st.columns(2)
                    with c_ws1:
                        working_set_pct = st.slider("Tỷ lệ Dữ liệu nóng (Working Set)", 10, 100, mongo_so_lieu.get("working_set", 20), 5)
                    with c_ws2:
                        index_overhead = st.slider("Overhead cho Index & Oplog", 10, 50, mongo_so_lieu.get("overhead", 15), 5)

                mongo_params = {
                    "method": "docs", "doc_count": doc_count, "doc_size_kb": avg_doc_size_kb,
                    "working_set": working_set_pct, "overhead": index_overhead,
                    "proof_key": uploaded_doc_proof, "proof_size": uploaded_size_proof
                }
            
            else: 
                with st.container(border=True):
                    st.warning("Lưu ý: Phương pháp này chỉ nhân bản cấu hình cũ.")
                    c1, c2, c3 = st.columns(3)
                    curr_ram = c1.number_input("RAM Hiện tại (GB)", value=32, key="m_ram")
                    curr_cpu = c2.number_input("vCPU Hiện tại", value=8, key="m_cpu")
                    curr_disk = c3.number_input("Disk Hiện tại (GB)", value=500, key="m_disk")
                    
                    l1, l2 = st.columns(2)
                    load_ram = l1.slider("% RAM đang dùng", 0, 100, 70, key="m_load_ram")
                    load_cpu = l2.slider("% CPU đang dùng", 0, 100, 30, key="m_load_cpu")
                    target_scale = st.number_input("Tỉ lệ mở rộng (Scale Ratio)", value=1.5, key="m_scale")

                    mongo_params = {
                        "method": "linear", "curr_ram": curr_ram, "curr_cpu": curr_cpu, "curr_disk": curr_disk, 
                        "load_ram": load_ram, "load_cpu": load_cpu, "scale": target_scale
                    }

    with col_result:
        st.subheader("Kết quả Định cỡ MongoDB")
        
        # Hiển thị dữ liệu đã lưu trước đó (nếu có)
        if mongo_result:
            with st.container(border=True):
                st.markdown("##### 📋 Kết quả đã lưu trước đó:")
                st.markdown(f"**Kiến trúc:** `{mongo_result.get('kien_truc', 'N/A')}`")
                st.markdown(f"**Số Node:** `{mongo_result.get('node', 'N/A')}`")
                
                col_m1, col_m2 = st.columns(2)
                col_m1.metric("vCPU", f"{mongo_result.get('vcpu', 0)} Core")
                col_m2.metric("RAM", f"{mongo_result.get('ram', 0)} GB")
                
                col_m3, col_m4 = st.columns(2)
                col_m3.metric("Disk", f"{mongo_result.get('disk', 0)} GB")
                col_m4.metric("IOPS", mongo_result.get('iops', 'N/A'))
        
        if st.button("TÍNH TOÁN TÀI NGUYÊN MONGODB", type="primary", use_container_width=True):
            
            final_model = ""
            node_config = {}
            total_storage_gb = 0
            
            if mongo_params["method"] == "docs":
                raw_data_gb = (mongo_params["doc_count"] * mongo_params["doc_size_kb"]) / (1024 * 1024)
                total_storage_gb = raw_data_gb * (1 + mongo_params["overhead"]/100) * growth_factor
                
                raw_ram_gb = total_storage_gb * (mongo_params["working_set"] / 100)
                req_ram = raw_ram_gb / ram_threshold
                
                req_cpu = req_ram / 4 
                if req_cpu < 4: req_cpu = 4 

                SHARDING_THRESHOLD = 1024 
                
                if total_storage_gb < SHARDING_THRESHOLD:
                    final_model = "Replica Set (3 Nodes)"
                    node_config = {"qty": 3, "vCPU": math.ceil(req_cpu), "RAM": math.ceil(req_ram), "Disk": math.ceil(total_storage_gb / disk_threshold)}
                    st.info(f"Dung lượng Data dự kiến: {total_storage_gb:.2f} GB (<1TB) -> Đề xuất mô hình **Replica Set** tiêu chuẩn.")
                else:
                    final_model = "Sharded Cluster"
                    num_shards = math.ceil(total_storage_gb / SHARDING_THRESHOLD)
                    if num_shards < 2: num_shards = 2
                    
                    total_data_nodes = num_shards * 3
                    
                    per_shard_storage = total_storage_gb / num_shards
                    per_shard_ram = req_ram / num_shards
                    per_shard_cpu = req_cpu / num_shards
                    
                    node_config = {"qty": total_data_nodes, "vCPU": max(4, math.ceil(per_shard_cpu)), "RAM": max(16, math.ceil(per_shard_ram)), "Disk": math.ceil(per_shard_storage / disk_threshold)}
                    st.info(f"Dung lượng Data dự kiến: {total_storage_gb:.2f} GB (>1TB) -> Đề xuất mô hình **Sharded Cluster** ({num_shards} Shards).")
                    st.warning(f"Lưu ý: Cần thêm 3 Node Config Server và ít nhất 2 Node Mongos (Router).")

            else: 
                p = mongo_params
                req_ram_total = (p["curr_ram"] * (p["load_ram"]/100) * p["scale"] * growth_factor) / ram_threshold
                req_cpu_total = (p["curr_cpu"] * (p["load_cpu"]/100) * p["scale"] * growth_factor) / cpu_threshold
                req_disk_total = (p["curr_disk"] * p["scale"] * growth_factor) / disk_threshold
                
                final_model = "Replica Set (Tuyến tính)"
                node_config = {"qty": 3, "vCPU": max(4, math.ceil(req_cpu_total)), "RAM": math.ceil(req_ram_total), "Disk": math.ceil(req_disk_total)}

            # Lưu kết quả vào session_state
            st.session_state.mongo_result = {
                "final_model": final_model,
                "node_config": node_config,
                "mongo_params": mongo_params,
                "mongo_desc": mongo_desc,
                "mongo_purpose": mongo_purpose,
                "uploaded_logic_img": uploaded_logic_img
            }

            with st.container(border=True):
                st.markdown(f"**Kiến trúc:** `{final_model}`")
                st.markdown(f"**Tổng Data Node:** `{node_config['qty']}`")
                
                st.divider()
                st.markdown("#### Cấu hình mỗi Node (Data Node)")
                
                m1, m2 = st.columns(2)
                m1.metric("vCPU", f"{node_config['vCPU']} Core")
                m2.metric("RAM", f"{node_config['RAM']} GB")
                
                m3, m4 = st.columns(2)
                m3.metric("Lưu trữ (Disk)", f"{node_config['Disk']} GB")
                m4.metric("IOPS Dự kiến", "3000+")

            st.divider()
            doc = Document()
            doc.add_heading('BÁO CÁO ĐỊNH CỠ MONGODB', 0)
            doc.add_paragraph(f"Mô hình: {final_model}")
            doc.add_paragraph(f"Số lượng Node: {node_config['qty']}")
            doc.add_paragraph(f"Cấu hình mỗi Node: {node_config['vCPU']} vCPU, {node_config['RAM']} GB RAM, {node_config['Disk']} GB Disk")
            
            b = BytesIO()
            doc.save(b)
            b.seek(0)
            st.download_button("Tải báo cáo (.docx)", data=b, file_name="mongodb_sizing.docx", use_container_width=True)

        # Nút lưu đặt bên ngoài if button để luôn hiển thị khi có kết quả
        if "mongo_result" in st.session_state:
            result = st.session_state.mongo_result
            
            if st.button("💾 Lưu kết quả MongoDB vào Database", use_container_width=True, key="save_mongo", type="primary"):
                if project_id:
                    mongo_params = result["mongo_params"]
                    node_config = result["node_config"]
                    final_model = result["final_model"]
                    
                    # Chuẩn bị dữ liệu theo cấu trúc mới
                    mongo_module_data = {
                        "dauVao": {
                            "logic_image": get_image_base64(result.get("uploaded_logic_img")),
                            "mo_ta": result.get("mongo_desc", ""),
                            "muc_dich": result.get("mongo_purpose", "")
                        },
                        "phuong_phap": mongo_params.get("method", ""),
                        "so_lieu": []
                    }
                    
                    # Thêm số liệu theo phương pháp
                    if mongo_params.get("method") == "docs":
                        mongo_module_data["so_lieu"].append({
                            "document": mongo_params.get("doc_count", 0),
                            "kich_thuoc": mongo_params.get("doc_size_kb", 0),
                            "document_image": get_image_base64(mongo_params.get("proof_key")),
                            "kich_thuoc_image": get_image_base64(mongo_params.get("proof_size")),
                            "working_set": mongo_params.get("working_set", 0),
                            "overhead": mongo_params.get("overhead", 0)
                        })
                    else:
                        mongo_module_data["so_lieu"].append({
                            "ram": mongo_params.get("curr_ram", 0),
                            "ram_percent": mongo_params.get("load_ram", 0),
                            "vcpu": mongo_params.get("curr_cpu", 0),
                            "vcpu_percent": mongo_params.get("load_cpu", 0),
                            "disk": mongo_params.get("curr_disk", 0),
                            "scale_ratio": mongo_params.get("scale", 1.0)
                        })
                    
                    # Thêm kết quả
                    mongo_module_data["result"] = {
                        "kien_truc": final_model,
                        "node": node_config.get("qty", 0),
                        "vcpu": node_config.get("vCPU", 0),
                        "ram": node_config.get("RAM", 0),
                        "disk": node_config.get("Disk", 0),
                        "iops": "3000+"
                    }
                    
                    success, msg = save_dinh_co_module(project_id, "mongo_module", mongo_module_data)
                    if success:
                        st.success("✅ Lưu dữ liệu MongoDB thành công!")
                    else:
                        st.error(f"❌ Lỗi: {msg}")
                else:
                    st.error("Chưa nhập Mã dự án (Project ID)")

# --- TRƯỜNG HỢP 3: POSTGRESQL SIZING ---
elif db_type == "PostgreSQL":
    
    # Load saved data
    pg_saved = st.session_state.get("pg_saved_data", {})
    pg_dauVao = pg_saved.get("dauVao", {})
    pg_saved_result = pg_saved.get("result", {})
    pg_so_lieu = pg_saved.get("so_lieu", [{}])[0] if pg_saved.get("so_lieu") else {}
    
    col_input, col_result = st.columns([1.2, 0.8], gap="medium")
    
    with col_input:
        st.subheader("Tham số Đầu vào (PostgreSQL)")
        
        # Tabs nhập liệu
        tab_info, tab_sizing = st.tabs(["Thông tin Mô hình", "Số liệu Tính toán"])
        
        with tab_info:
            with st.container(border=True):
                st.markdown("**1. Sơ đồ Logic HA**")
                
                col_upload, col_icon = st.columns([5, 1])
                
                with col_upload:
                    uploaded_logic_img = st.file_uploader("Tải lên sơ đồ HA", type=['png', 'jpg'], key="pg_logic", label_visibility="collapsed")
                
                with col_icon:
                    with st.popover("❓", help="Xem ảnh minh họa"):
                        st.image("mo-hinh-he-thong/postgresql_example.png", use_container_width=True, caption="Ví dụ sơ đồ Logic PostgreSQL")
                
                if uploaded_logic_img:
                    st.image(uploaded_logic_img, use_container_width=True, caption="Sơ đồ Logic HA")

            with st.container(border=True):
                st.markdown("**3. Thông tin Nghiệp vụ**")
                pg_desc = st.text_area("Mô tả module", height=100, value=pg_dauVao.get("mo_ta", ""), placeholder="Ví dụ: Core Banking, Payment Gateway...", key="pg_desc")
                pg_purpose = st.text_area("Mục đích sử dụng", height=100, value=pg_dauVao.get("muc_dich", ""), placeholder="Ví dụ: OLTP Transaction, Data Warehousing...", key="pg_purp")

        with tab_sizing:
            method = st.radio("Phương pháp tính toán", ["Theo Dung lượng & TPS (Khuyên dùng)", "Tuyến tính theo cấu hình cũ"], horizontal=True, key="pg_method")
            
            pg_params = {}

            if "Dung lượng" in method:
                with st.container(border=True):
                    st.markdown("#### Dự báo Dữ liệu (Storage)")
                    c_st1, c_st2 = st.columns(2)
                    
                    with c_st1:
                        raw_data_gb = st.number_input("Dung lượng Dữ liệu thô (GB)", min_value=10, value=pg_so_lieu.get("dung_luong", 500), step=50, help="Dung lượng thực tế chưa bao gồm Index")
                        st.caption("Dữ liệu dự kiến trong 1-2 năm tới")
                    
                    with c_st2:
                        index_ratio = st.slider("Tỷ lệ Index & Bloat (%)", 20, 100, pg_so_lieu.get("index_ratio", 50), help="Postgres B-Tree Index khá nặng + cơ chế MVCC gây bloat")
                        st.caption("Khuyến nghị: 40-50% cho hệ thống OLTP")

                with st.container(border=True):
                    st.markdown("#### Hiệu năng (Compute)")
                    c_perf1, c_perf2 = st.columns(2)
                    
                    with c_perf1:
                        tps_target = st.number_input("TPS Mục tiêu (Transaction/s)", value=pg_so_lieu.get("tps_target", 1000), step=100, help="Số lượng giao dịch ghi/đọc mỗi giây")
                    
                    with c_perf2:
                        # Postgres dựa nhiều vào OS Cache
                        working_set = st.slider("Tỷ lệ Working Set (%)", 10, 80, pg_so_lieu.get("working_set", 25), help="% Dữ liệu thường xuyên truy cập cần nằm trên RAM")

                # Cấu hình HA
                with st.container(border=True):
                    st.markdown("#### Mô hình High Availability (HA)")
                    ha_mode = st.selectbox("Kiến trúc HA", [
                        "Patroni Cluster (3 Nodes - Khuyên dùng)", 
                        "Primary - Standby (2 Nodes)",
                        "Standalone (1 Node - Dev/Test)"
                    ])

                pg_params = {
                    "method": "calc", 
                    "raw_data": raw_data_gb, 
                    "index_ratio": index_ratio,
                    "tps": tps_target,
                    "working_set": working_set,
                    "ha_mode": ha_mode
                }
            
            else: # Tuyến tính
                with st.container(border=True):
                    st.warning("Lưu ý: Phương pháp này chỉ nhân bản cấu hình cũ.")
                    c1, c2, c3 = st.columns(3)
                    curr_ram = c1.number_input("RAM Hiện tại (GB)", value=32, key="pg_ram")
                    curr_cpu = c2.number_input("vCPU Hiện tại", value=8, key="pg_cpu")
                    curr_disk = c3.number_input("Disk Hiện tại (GB)", value=500, key="pg_disk")
                    
                    l1, l2 = st.columns(2)
                    load_ram = l1.slider("% RAM đang dùng", 0, 100, 70, key="pg_load_ram")
                    load_cpu = l2.slider("% CPU đang dùng", 0, 100, 30, key="pg_load_cpu")
                    target_scale = st.number_input("Tỉ lệ mở rộng (Scale Ratio)", value=1.5, key="pg_scale")

                    pg_params = {
                        "method": "linear", 
                        "curr_ram": curr_ram, "curr_cpu": curr_cpu, "curr_disk": curr_disk, 
                        "load_ram": load_ram, "load_cpu": load_cpu, 
                        "scale": target_scale
                    }

    # Cột hiển thị kết quả
    with col_result:
        st.subheader("Kết quả Định cỡ PostgreSQL")
        
        # Hiển thị dữ liệu đã lưu trước đó (nếu có)
        if pg_saved_result:
            with st.container(border=True):
                st.markdown("##### 📋 Kết quả đã lưu trước đó:")
                st.markdown(f"**Kiến trúc:** `{pg_saved_result.get('kien_truc', 'N/A')}`")
                st.markdown(f"**Số Node:** `{pg_saved_result.get('node', 'N/A')}`")
                
                col_p1, col_p2 = st.columns(2)
                col_p1.metric("vCPU", f"{pg_saved_result.get('vcpu', 0)} Core")
                col_p2.metric("RAM", f"{pg_saved_result.get('ram', 0)} GB")
                
                col_p3, col_p4 = st.columns(2)
                col_p3.metric("Disk", f"{pg_saved_result.get('disk', 0)} GB")
                col_p4.metric("IOPS", pg_saved_result.get('iops', 'N/A'))
        
        if st.button("TÍNH TOÁN TÀI NGUYÊN POSTGRES", type="primary", use_container_width=True):
            
            final_model = ""
            node_config = {}
            note_storage = ""
            
            # --- LOGIC TÍNH TOÁN POSTGRES ---
            if pg_params["method"] == "calc":
                # 1. Tính Storage (Mỗi node phải chứa Full Data)
                # Total Disk = Raw * (1 + Index%) * Growth
                disk_need = pg_params["raw_data"] * (1 + pg_params["index_ratio"]/100) * growth_factor
                disk_final = math.ceil(disk_need / disk_threshold)
                
                # 2. Tính RAM
                # RAM Need = (Total Disk Used * Working Set) / Ram Threshold
                # Working Set ở đây tính trên Data thực (Raw + Index)
                data_on_disk = pg_params["raw_data"] * (1 + pg_params["index_ratio"]/100)
                ram_need = (data_on_disk * (pg_params["working_set"]/100)) / ram_threshold
                # Ràng buộc tối thiểu cho OS và Shared Buffers
                if ram_need < 8: ram_need = 8
                
                # 3. Tính CPU (Dựa trên TPS và RAM)
                # Rule of thumb: 1 vCPU gánh được khoảng 200-500 TPS tùy complexity. Lấy trung bình 300.
                cpu_by_tps = pg_params["tps"] / 300 
                # Rule of thumb: 1 vCPU cho 4GB RAM (để quản lý buffer, connections)
                cpu_by_ram = ram_need / 4
                
                cpu_final = max(cpu_by_tps, cpu_by_ram)
                if cpu_final < 4: cpu_final = 4

                # 4. Mô hình
                final_model = pg_params["ha_mode"]
                if "3 Nodes" in final_model:
                    qty = 3
                    note_storage = "Dữ liệu được NHÂN BẢN (Duplicate) trên cả 3 Node."
                elif "2 Nodes" in final_model:
                    qty = 2
                    note_storage = "Dữ liệu được NHÂN BẢN (Duplicate) trên cả 2 Node."
                else:
                    qty = 1
                    note_storage = "Chạy đơn lẻ (Không có HA)."

                node_config = {
                    "qty": qty, 
                    "vCPU": math.ceil(cpu_final), 
                    "RAM": math.ceil(ram_need), 
                    "Disk": disk_final
                }
                
                st.info(f"Dung lượng lưu trữ yêu cầu (bao gồm Index & Bloat): **{disk_need:.2f} GB**")

            else: # Tuyến tính
                p = pg_params
                req_ram_total = (p["curr_ram"] * (p["load_ram"]/100) * p["scale"] * growth_factor) / ram_threshold
                req_cpu_total = (p["curr_cpu"] * (p["load_cpu"]/100) * p["scale"] * growth_factor) / cpu_threshold
                req_disk_total = (p["curr_disk"] * p["scale"] * growth_factor) / disk_threshold
                disk_need = p["curr_disk"]  # Gán giá trị mặc định
                
                final_model = "Patroni HA (Tuyến tính)"
                node_config = {
                    "qty": 3, # Mặc định HA 3 node
                    "vCPU": max(4, math.ceil(req_cpu_total)), 
                    "RAM": math.ceil(req_ram_total), 
                    "Disk": math.ceil(req_disk_total)
                }
                note_storage = "Tính theo cấu hình hiện tại nhân bản."

            # Lưu kết quả vào session_state
            st.session_state.pg_result = {
                "final_model": final_model,
                "node_config": node_config,
                "pg_params": pg_params,
                "pg_desc": pg_desc,
                "pg_purpose": pg_purpose,
                "uploaded_logic_img": uploaded_logic_img,
                "disk_need": disk_need,
                "note_storage": note_storage
            }

            # Hiển thị Kết quả (Card Style)
            with st.container(border=True):
                st.markdown(f"**Kiến trúc:** `{final_model}`")
                st.markdown(f"**Số lượng Server:** `{node_config['qty']}`")
                
                st.divider()
                st.markdown("#### Cấu hình mỗi Node")
                
                m1, m2 = st.columns(2)
                m1.metric("vCPU", f"{node_config['vCPU']} Core", delta_color="off")
                m2.metric("RAM", f"{node_config['RAM']} GB", help="Shared Buffers + OS Cache")
                
                m3, m4 = st.columns(2)
                m3.metric("Lưu trữ (Disk)", f"{node_config['Disk']} GB", "Local SSD")
                m4.metric("Dữ liệu", "Duplicate", help="Mỗi Node chứa 100% dữ liệu")

            st.caption(f"ℹ️ *Ghi chú Storage: {note_storage}*")

            # Khu vực Xuất báo cáo
            st.divider()
            doc = Document()
            doc.add_heading('BÁO CÁO ĐỊNH CỠ POSTGRESQL', 0)
            doc.add_paragraph(f"Mô hình: {final_model}")
            doc.add_paragraph(f"Số lượng Node: {node_config['qty']}")
            doc.add_paragraph(f"Cấu hình mỗi Node: {node_config['vCPU']} vCPU, {node_config['RAM']} GB RAM, {node_config['Disk']} GB Disk")
            doc.add_paragraph(f"Lưu ý: {note_storage}")
            
            b = BytesIO()
            doc.save(b)
            b.seek(0)
            st.download_button("Tải báo cáo (.docx)", data=b, file_name="postgres_sizing.docx", use_container_width=True)

        # Nút lưu đặt bên ngoài if button để luôn hiển thị khi có kết quả
        if "pg_result" in st.session_state:
            result = st.session_state.pg_result
            
            if st.button("💾 Lưu kết quả PostgreSQL vào Database", use_container_width=True, key="save_pg", type="primary"):
                if project_id:
                    pg_params = result["pg_params"]
                    node_config = result["node_config"]
                    final_model = result["final_model"]
                    disk_need = result["disk_need"]
                    note_storage = result["note_storage"]
                    
                    # Chuẩn bị dữ liệu theo cấu trúc mới
                    pg_module_data = {
                        "dauVao": {
                            "logic_image": get_image_base64(result.get("uploaded_logic_img")),
                            "mo_ta": result.get("pg_desc", ""),
                            "muc_dich": result.get("pg_purpose", "")
                        },
                        "phuong_phap": pg_params.get("method", ""),
                        "so_lieu": []
                    }
                    
                    # Thêm số liệu theo phương pháp
                    if pg_params.get("method") == "calc":
                        pg_module_data["so_lieu"].append({
                            "storage": pg_params.get("raw_data", 0),
                            "storage_percent": pg_params.get("index_ratio", 0),
                            "compute": pg_params.get("tps", 0),
                            "compute_percent": pg_params.get("working_set", 0),
                            "ha": pg_params.get("ha_mode", "")
                        })
                    else:
                        pg_module_data["so_lieu"].append({
                            "ram": pg_params.get("curr_ram", 0),
                            "ram_percent": pg_params.get("load_ram", 0),
                            "vcpu": pg_params.get("curr_cpu", 0),
                            "vcpu_percent": pg_params.get("load_cpu", 0),
                            "disk": pg_params.get("curr_disk", 0),
                            "scale_ratio": pg_params.get("scale", 1.0)
                        })
                    
                    # Thêm kết quả
                    pg_module_data["result"] = {
                        "storage_require": disk_need,
                        "kien_truc": final_model,
                        "node": node_config.get("qty", 0),
                        "vcpu": node_config.get("vCPU", 0),
                        "ram": node_config.get("RAM", 0),
                        "disk": node_config.get("Disk", 0),
                        "du_lieu": note_storage
                    }
                    
                    success, msg = save_dinh_co_module(project_id, "postgresql_module", pg_module_data)
                    if success:
                        st.success("✅ Lưu dữ liệu PostgreSQL thành công!")
                    else:
                        st.error(f"❌ Lỗi: {msg}")
                else:
                    st.error("Chưa nhập Mã dự án (Project ID)")

# --- TRƯỜNG HỢP 4: ORACLE/MARIA/DB THƯỜNG ---
else:
    
    col_input, col_result = st.columns([1.2, 0.8], gap="large")
    
    # Load dữ liệu đã lưu nếu có
    oracle_saved = st.session_state.get("oracle_saved_data", {})
    baseline_saved = oracle_saved.get("baseline", {})
    storage_saved = oracle_saved.get("storage", [{}])[0] if oracle_saved.get("storage") else {}
    result_saved = oracle_saved.get("result", {})
    
    with col_input:
        st.subheader("A. Hệ thống Tham chiếu (Baseline)")
        
        with st.container(border=True):
            c1, c2 = st.columns(2)
            current_ccu = c1.number_input("CCU Hiện tại", value=baseline_saved.get("ccu_hien_tai", 100))
            target_ccu = c2.number_input("CCU Mục tiêu", value=baseline_saved.get("ccu_muc_tieu", 200))
            scale_ratio = target_ccu / current_ccu
            st.caption(f"Tỉ lệ Scale: **{scale_ratio:.2f}x**")

        st.markdown("#### Danh sách Server hiện tại")
        
        for i, server in enumerate(st.session_state.servers):
            with st.container(border=True):
                head_col1, head_col2 = st.columns([0.85, 0.15])
                head_col1.markdown(f"**Server {i + 1}**")
                if head_col2.button("🗑️", key=f"btn_del_{i}"):
                    remove_server(i)
                    st.rerun()

                row1 = st.columns([2, 1, 1])
                server["ip"] = row1[0].text_input(f"IP Sv{i+1}", value=server["ip"], key=f"ip_{i}")
                server["vcpu"] = row1[1].number_input(f"vCPU Sv{i+1}", value=server["vcpu"], key=f"cpu_{i}")
                server["ram"] = row1[2].number_input(f"RAM(GB) Sv{i+1}", value=server["ram"], key=f"ram_{i}")

                row2 = st.columns(2)
                server["cpu_load"] = row2[0].slider(f"% CPU Used Sv{i+1}", 0, 100, server["cpu_load"], key=f"sld_cpu_{i}")
                server["ram_load"] = row2[1].slider(f"% RAM Used Sv{i+1}", 0, 100, server["ram_load"], key=f"sld_ram_{i}")

        btn_col1, btn_col2 = st.columns([1, 1])
        if btn_col1.button("➕ Thêm Server", type="secondary", use_container_width=True):
            add_server()
            st.rerun()
            
        st.markdown("---")
        st.markdown("#### Thông tin Lưu trữ (Storage)")
        with st.container(border=True):
            sd1, sd2 = st.columns(2)
            current_storage_used_data = sd1.number_input("Data Used (GB)", value=storage_saved.get("data_used", 500))
            current_storage_used_log = sd2.number_input("Log Used (GB)", value=storage_saved.get("log_used", 100))
            current_storage_used_backup = st.number_input("Backup Full (GB)", value=storage_saved.get("backup_used", 500))

    with col_result:
        st.subheader("B. Kết quả & Đề xuất")
        
        # Hiển thị kết quả đã lưu nếu có
        if result_saved:
            st.info("📋 Đã tìm thấy kết quả đã lưu trước đó")
            with st.container(border=True):
                st.markdown(f"**Kết quả đã lưu:**")
                st.markdown(f"- Số Node: {result_saved.get('node', 'N/A')}")
                st.markdown(f"- vCPU/Node: {result_saved.get('vCpu', 'N/A')} Core")
                st.markdown(f"- RAM/Node: {result_saved.get('ram', 'N/A')} GB")
                st.markdown(f"- Data Volume: {result_saved.get('data_volume', 'N/A')} GB")
                st.markdown(f"- Log Volume: {result_saved.get('log_volume', 'N/A')} GB")
            st.divider()
        
        with st.container(border=True):
            st.markdown("##### Cấu hình Mục tiêu")
            min_node = 2 if db_type == "Oracle RAC" else 1
            proposed_n = st.number_input("Số Node đề xuất:", min_value=min_node, value=max(len(st.session_state.servers), min_node))
            
            calc_btn = st.button("TÍNH TOÁN SIZING", type="primary", use_container_width=True)

        if calc_btn:
            total_cpu_used = 0
            total_ram_used = 0
            
            for s in st.session_state.servers:
                total_cpu_used += s["vcpu"] * (s["cpu_load"] / 100)
                total_ram_used += s["ram"] * (s["ram_load"] / 100)

            req_cpu = (total_cpu_used * scale_ratio * growth_factor) / cpu_threshold
            req_ram = (total_ram_used * scale_ratio * growth_factor) / ram_threshold
            
            req_disk_data = (current_storage_used_data * scale_ratio * growth_factor) / disk_threshold
            req_disk_log = (current_storage_used_log * scale_ratio * growth_factor) / disk_threshold
            
            per_node_cpu = math.ceil(math.ceil(req_cpu) / proposed_n)
            per_node_ram = math.ceil(math.ceil(req_ram) / proposed_n)
            
            is_shared = db_type == "Oracle RAC"
            storage_txt = f"{math.ceil(req_disk_data)} GB (Shared)" if is_shared else f"{math.ceil(req_disk_data)} GB (Mỗi Node)"

            # Lưu kết quả vào session_state để dùng khi bấm nút Lưu
            st.session_state.oracle_result = {
                "current_ccu": current_ccu,
                "target_ccu": target_ccu,
                "proposed_n": proposed_n,
                "per_node_cpu": per_node_cpu,
                "per_node_ram": per_node_ram,
                "req_disk_data": req_disk_data,
                "req_disk_log": req_disk_log,
                "storage_data": current_storage_used_data,
                "storage_log": current_storage_used_log,
                "storage_backup": current_storage_used_backup,
                "db_type": db_type
            }

            st.success(f"✅ Kết quả tính toán cho {db_type} ({proposed_n} Node)")
            
            with st.container(border=True):
                st.markdown("#### Cấu hình Mỗi Node")
                m1, m2 = st.columns(2)
                m1.metric("vCPU", f"{per_node_cpu} Core")
                m2.metric("RAM", f"{per_node_ram} GB")
                
                st.divider()
                st.markdown("#### Dung lượng Lưu trữ")
                m3, m4 = st.columns(2)
                m3.metric("Data Volume", storage_txt)
                m4.metric("Log Volume", f"{math.ceil(req_disk_log/proposed_n)} GB")

            doc = Document()
            doc.add_heading(f'BÁO CÁO SIZING: {db_type.upper()}', 0)
            
            doc.add_heading('1. Thông tin Baseline', level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Server IP'
            hdr[1].text = 'vCPU (Used)'
            hdr[2].text = 'RAM (Used)'
            
            for s in st.session_state.servers:
                row = table.add_row().cells
                row[0].text = s['ip']
                row[1].text = f"{s['vcpu']} ({s['cpu_load']}%)"
                row[2].text = f"{s['ram']} ({s['ram_load']}%)"
            
            doc.add_heading('2. Kết quả Đề xuất', level=1)
            doc.add_paragraph(f"Tổng số Node: {proposed_n}")
            doc.add_paragraph(f"Cấu hình mỗi Node: {per_node_cpu} vCPU, {per_node_ram} GB RAM")
            
            b = BytesIO()
            doc.save(b)
            b.seek(0)
            
            c_down, c_save = st.columns(2)
            with c_down:
                st.download_button("📥 Tải báo cáo (.docx)", data=b, file_name=f"sizing_{db_type}.docx", use_container_width=True)
            
        # Nút lưu đặt bên ngoài if calc_btn để luôn hiển thị khi có kết quả
        if "oracle_result" in st.session_state:
            result = st.session_state.oracle_result
            
            if st.button("💾 Lưu kết quả vào Database", use_container_width=True, key="save_oracle", type="primary"):
                if project_id:
                    # Chuẩn bị dữ liệu theo cấu trúc mới
                    oracle_module_data = {
                        "baseline": {
                            "ccu_hien_tai": result["current_ccu"],
                            "ccu_muc_tieu": result["target_ccu"]
                        },
                        "server": [],
                        "storage": [{
                            "data_used": result["storage_data"],
                            "log_used": result["storage_log"],
                            "backup_used": result["storage_backup"]
                        }],
                        "result": {
                            "node": result["proposed_n"],
                            "vCpu": result["per_node_cpu"],
                            "ram": result["per_node_ram"],
                            "data_volume": math.ceil(result["req_disk_data"]),
                            "log_volume": math.ceil(result["req_disk_log"] / result["proposed_n"])
                        }
                    }
                    
                    # Thêm thông tin các server
                    for s in st.session_state.servers:
                        oracle_module_data["server"].append({
                            "ip": s["ip"],
                            "vCpu": s["vcpu"],
                            "ram": s["ram"],
                            "vCpu_percent": s["cpu_load"],
                            "ram_percent": s["ram_load"]
                        })
                    
                    success, msg = save_dinh_co_module(project_id, "oracleRac_module", oracle_module_data)
                    if success:
                        st.success("✅ Lưu dữ liệu Oracle RAC thành công!")
                    else:
                        st.error(f"❌ Lỗi: {msg}")
                else:
                    st.error("Chưa nhập Mã dự án (Project ID)")